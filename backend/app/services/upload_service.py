"""
Servicio para manejo de archivos subidos
"""
import os
import tempfile
import hashlib
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime
import PyPDF2
import docx
import pandas as pd
from ..core.utils import sanitize_filename, chunk_text, timing_decorator
from ..vectorstore import agregar_documento
from ..config import UPLOAD_DIR, MAX_FILE_SIZE

class UploadService:
    """Servicio para manejo de archivos subidos"""
    
    def __init__(self):
        self.supported_extensions = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.xlsx': self._extract_excel_data,
            '.xls': self._extract_excel_data,
            '.csv': self._extract_csv_data,
            '.txt': self._extract_txt_text
        }
        self.max_file_size = MAX_FILE_SIZE
        self.upload_dir = Path(UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
    
    @timing_decorator
    async def process_file_upload(self, file_content: bytes, filename: str, 
                                user_id: str = "anonymous") -> Dict:
        """Procesa un archivo subido"""
        try:
            # 1. Validaciones iniciales
            validation_result = self._validate_file(file_content, filename)
            if not validation_result["is_valid"]:
                return {
                    "success": False,
                    "error": validation_result["error"],
                    "file_info": None
                }
            
            # 2. Generar información del archivo
            file_info = self._generate_file_info(file_content, filename, user_id)
            
            # 3. Verificar si ya existe (por hash)
            existing_file = self._check_existing_file(file_info["hash"])
            if existing_file:
                return {
                    "success": True,
                    "message": "Archivo ya procesado anteriormente",
                    "file_info": existing_file,
                    "is_duplicate": True
                }
            
            # 4. Guardar archivo físicamente
            file_path = await self._save_file(file_content, file_info)
            
            # 5. Extraer contenido según tipo
            extracted_content = await self._extract_content(file_path, filename)
            
            # 6. Procesar y vectorizar contenido
            processing_result = await self._process_content(extracted_content, file_info)
            
            # 7. Guardar metadata
            await self._save_file_metadata(file_info, processing_result)
            
            return {
                "success": True,
                "message": "Archivo procesado exitosamente",
                "file_info": file_info,
                "processing_stats": processing_result,
                "is_duplicate": False
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Error procesando archivo: {str(e)}",
                "file_info": None
            }
    
    def _validate_file(self, file_content: bytes, filename: str) -> Dict:
        """Valida archivo antes de procesar"""
        # Verificar tamaño
        if len(file_content) > self.max_file_size:
            return {
                "is_valid": False,
                "error": f"Archivo muy grande. Máximo permitido: {self.max_file_size/1024/1024:.1f}MB"
            }
        
        # Verificar extensión
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.supported_extensions:
            return {
                "is_valid": False,
                "error": f"Tipo de archivo no soportado. Extensiones permitidas: {', '.join(self.supported_extensions.keys())}"
            }
        
        # Verificar que no esté vacío
        if len(file_content) == 0:
            return {
                "is_valid": False,
                "error": "El archivo está vacío"
            }
        
        return {"is_valid": True, "error": None}
    
    def _generate_file_info(self, file_content: bytes, filename: str, user_id: str) -> Dict:
        """Genera información del archivo"""
        # Calcular hash para identificación única
        file_hash = hashlib.md5(file_content).hexdigest()
        
        # Sanitizar nombre
        safe_filename = sanitize_filename(filename)
        
        # Generar nombre único
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_filename = f"{timestamp}_{file_hash[:8]}_{safe_filename}"
        
        return {
            "original_name": filename,
            "safe_name": safe_filename,
            "unique_name": unique_filename,
            "hash": file_hash,
            "size": len(file_content),
            "extension": Path(filename).suffix.lower(),
            "upload_time": datetime.now().isoformat(),
            "user_id": user_id,
            "status": "processing"
        }
    
    def _check_existing_file(self, file_hash: str) -> Optional[Dict]:
        """Verifica si el archivo ya fue procesado"""
        # En implementación real, consultaría base de datos
        # Por ahora retorna None (archivo nuevo)
        return None
    
    async def _save_file(self, file_content: bytes, file_info: Dict) -> Path:
        """Guarda archivo en disco"""
        file_path = self.upload_dir / file_info["unique_name"]
        
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        return file_path
    
    async def _extract_content(self, file_path: Path, original_filename: str) -> Dict:
        """Extrae contenido según tipo de archivo"""
        file_ext = file_path.suffix.lower()
        
        if file_ext in self.supported_extensions:
            extractor = self.supported_extensions[file_ext]
            return await extractor(file_path)
        else:
            raise ValueError(f"Tipo de archivo no soportado: {file_ext}")
    
    # EXTRACTORES ESPECÍFICOS POR TIPO DE ARCHIVO
    
    async def _extract_pdf_text(self, file_path: Path) -> Dict:
        """Extrae texto de archivo PDF"""
        try:
            text_content = ""
            metadata = {"pages": 0, "has_images": False}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n--- Página {page_num + 1} ---\n"
                            text_content += page_text + "\n"
                    except Exception as e:
                        print(f"Error extrayendo página {page_num + 1}: {e}")
                        continue
            
            return {
                "type": "text",
                "content": text_content.strip(),
                "metadata": metadata,
                "extraction_method": "PyPDF2"
            }
            
        except Exception as e:
            raise Exception(f"Error extrayendo PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_path: Path) -> Dict:
        """Extrae texto de archivo DOCX"""
        try:
            doc = docx.Document(file_path)
            
            text_content = ""
            metadata = {"paragraphs": 0, "tables": 0}
            
            # Extraer párrafos
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content += para.text + "\n"
                    metadata["paragraphs"] += 1
            
            # Extraer tablas
            for table in doc.tables:
                metadata["tables"] += 1
                text_content += "\n--- Tabla ---\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    text_content += row_text + "\n"
            
            return {
                "type": "text",
                "content": text_content.strip(),
                "metadata": metadata,
                "extraction_method": "python-docx"
            }
            
        except Exception as e:
            raise Exception(f"Error extrayendo DOCX: {str(e)}")
    
    async def _extract_excel_data(self, file_path: Path) -> Dict:
        """Extrae datos de archivo Excel"""
        try:
            # Leer todas las hojas
            excel_data = pd.read_excel(file_path, sheet_name=None)
            
            text_content = ""
            metadata = {"sheets": len(excel_data), "total_rows": 0, "total_cols": 0}
            
            for sheet_name, df in excel_data.items():
                text_content += f"\n--- Hoja: {sheet_name} ---\n"
                
                # Convertir DataFrame a texto estructurado
                if not df.empty:
                    # Encabezados
                    headers = " | ".join([str(col) for col in df.columns])
                    text_content += f"Columnas: {headers}\n\n"
                    
                    # Datos (máximo 100 filas para evitar sobrecarga)
                    sample_df = df.head(100)
                    for index, row in sample_df.iterrows():
                        row_text = " | ".join([str(val) for val in row.values])
                        text_content += f"{row_text}\n"
                    
                    if len(df) > 100:
                        text_content += f"\n... (mostrando 100 de {len(df)} filas) ...\n"
                    
                    metadata["total_rows"] += len(df)
                    metadata["total_cols"] = max(metadata["total_cols"], len(df.columns))
                
                text_content += "\n"
            
            return {
                "type": "structured_data",
                "content": text_content.strip(),
                "metadata": metadata,
                "extraction_method": "pandas",
                "raw_data": excel_data  # Guardar datos originales para análisis
            }
            
        except Exception as e:
            raise Exception(f"Error extrayendo Excel: {str(e)}")
    
    async def _extract_csv_data(self, file_path: Path) -> Dict:
        """Extrae datos de archivo CSV"""
        try:
            # Intentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise Exception("No se pudo decodificar el archivo CSV")
            
            text_content = ""
            metadata = {"rows": len(df), "columns": len(df.columns)}
            
            # Encabezados
            headers = " | ".join([str(col) for col in df.columns])
            text_content += f"Columnas: {headers}\n\n"
            
            # Datos (máximo 100 filas)
            sample_df = df.head(100)
            for index, row in sample_df.iterrows():
                row_text = " | ".join([str(val) for val in row.values])
                text_content += f"{row_text}\n"
            
            if len(df) > 100:
                text_content += f"\n... (mostrando 100 de {len(df)} filas) ...\n"
            
            return {
                "type": "structured_data",
                "content": text_content.strip(),
                "metadata": metadata,
                "extraction_method": "pandas",
                "raw_data": df
            }
            
        except Exception as e:
            raise Exception(f"Error extrayendo CSV: {str(e)}")
    
    async def _extract_txt_text(self, file_path: Path) -> Dict:
        """Extrae texto de archivo TXT"""
        try:
            # Intentar diferentes encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise Exception("No se pudo decodificar el archivo TXT")
            
            metadata = {
                "lines": len(content.split('\n')),
                "characters": len(content),
                "words": len(content.split())
            }
            
            return {
                "type": "text",
                "content": content.strip(),
                "metadata": metadata,
                "extraction_method": "text_file"
            }
            
        except Exception as e:
            raise Exception(f"Error extrayendo TXT: {str(e)}")
    
    async def _process_content(self, extracted_content: Dict, file_info: Dict) -> Dict:
        """Procesa el contenido extraído para vectorización"""
        try:
            content = extracted_content["content"]
            
            if not content or len(content.strip()) < 50:
                return {
                    "chunks_created": 0,
                    "total_characters": 0,
                    "processing_status": "skipped_too_short"
                }
            
            # Dividir en chunks
            chunks = chunk_text(content, max_length=1000, overlap=100)
            
            # Vectorizar y guardar chunks
            chunks_created = 0
            for i, chunk in enumerate(chunks):
                if len(chunk.strip()) > 50:  # Solo chunks significativos
                    try:
                        # Agregar al vectorstore
                        chunk_metadata = {
                            "source_file": file_info["original_name"],
                            "chunk_index": i,
                            "file_hash": file_info["hash"],
                            "upload_time": file_info["upload_time"],
                            "file_type": extracted_content["type"]
                        }
                        
                        await agregar_documento(chunk, chunk_metadata)
                        chunks_created += 1
                        
                    except Exception as e:
                        print(f"Error procesando chunk {i}: {e}")
                        continue
            
            return {
                "chunks_created": chunks_created,
                "total_chunks": len(chunks),
                "total_characters": len(content),
                "processing_status": "completed",
                "extraction_metadata": extracted_content.get("metadata", {})
            }
            
        except Exception as e:
            return {
                "chunks_created": 0,
                "total_characters": 0,
                "processing_status": "error",
                "error": str(e)
            }
    
    async def _save_file_metadata(self, file_info: Dict, processing_result: Dict) -> None:
        """Guarda metadata del archivo procesado"""
        try:
            # En implementación real, guardaría en base de datos
            metadata = {
                **file_info,
                **processing_result,
                "status": "completed" if processing_result["chunks_created"] > 0 else "failed"
            }
            
            # Por ahora, guardar en archivo JSON local
            metadata_file = self.upload_dir / f"{file_info['hash']}_metadata.json"
            import json
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2, ensure_ascii=False)
                
        except Exception as e:
            print(f"Error guardando metadata: {e}")
    
    def get_uploaded_files(self, user_id: str = None) -> List[Dict]:
        """Obtiene lista de archivos subidos"""
        try:
            files = []
            metadata_files = list(self.upload_dir.glob("*_metadata.json"))
            
            for metadata_file in metadata_files:
                try:
                    import json
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        metadata = json.load(f)
                    
                    if user_id is None or metadata.get("user_id") == user_id:
                        files.append({
                            "id": metadata["hash"],
                            "name": metadata["original_name"],
                            "size": metadata["size"],
                            "type": metadata["extension"],
                            "upload_time": metadata["upload_time"],
                            "status": metadata.get("status", "unknown"),
                            "chunks_created": metadata.get("chunks_created", 0)
                        })
                        
                except Exception as e:
                    print(f"Error leyendo metadata {metadata_file}: {e}")
                    continue
            
            # Ordenar por fecha de subida (más reciente primero)
            files.sort(key=lambda x: x["upload_time"], reverse=True)
            return files
            
        except Exception as e:
            print(f"Error obteniendo archivos: {e}")
            return []
    
    def delete_file(self, file_hash: str, user_id: str = None) -> Dict:
        """Elimina un archivo y su metadata"""
        try:
            # Buscar metadata
            metadata_file = self.upload_dir / f"{file_hash}_metadata.json"
            
            if not metadata_file.exists():
                return {"success": False, "error": "Archivo no encontrado"}
            
            # Verificar permisos de usuario
            if user_id:
                import json
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    metadata = json.load(f)
                
                if metadata.get("user_id") != user_id:
                    return {"success": False, "error": "No tienes permisos para eliminar este archivo"}
            
            # Eliminar archivo físico
            for file_path in self.upload_dir.glob(f"*{file_hash}*"):
                if file_path.is_file():
                    file_path.unlink()
            
            # TODO: Eliminar chunks del vectorstore
            
            return {"success": True, "message": "Archivo eliminado correctamente"}
            
        except Exception as e:
            return {"success": False, "error": f"Error eliminando archivo: {str(e)}"}
    
    def get_file_stats(self) -> Dict:
        """Obtiene estadísticas de archivos subidos"""
        try:
            files = self.get_uploaded_files()
            
            total_files = len(files)
            total_size = sum(f["size"] for f in files)
            
            by_type = {}
            by_status = {}
            total_chunks = 0
            
            for file in files:
                # Por tipo
                file_type = file["type"]
                by_type[file_type] = by_type.get(file_type, 0) + 1
                
                # Por status
                status = file["status"]
                by_status[status] = by_status.get(status, 0) + 1
                
                # Chunks
                total_chunks += file.get("chunks_created", 0)
            
            return {
                "total_files": total_files,
                "total_size_mb": round(total_size / 1024 / 1024, 2),
                "total_chunks": total_chunks,
                "by_type": by_type,
                "by_status": by_status,
                "avg_chunks_per_file": round(total_chunks / max(total_files, 1), 1)
            }
            
        except Exception as e:
            return {"error": f"Error obteniendo estadísticas: {str(e)}"}

# Instancia global del servicio
upload_service = UploadService()